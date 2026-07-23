"""Word 文档处理工具。

使用 python-docx 库提供 Word 文档的读取、写入、替换等操作。
所有方法均为静态方法，无状态。
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class WordTool:
    """Word 文档处理工具。静态方法，无状态。"""

    # ------------------------------------------------------------------
    # 内部辅助
    # ------------------------------------------------------------------

    @staticmethod
    def _resolve_path(filepath: str) -> Path:
        """解析并校验文件路径，支持中文路径。"""
        path = Path(filepath).resolve()
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {filepath}")
        if not path.suffix.lower() in (".docx",):
            raise ValueError(f"仅支持 .docx 格式，当前文件: {path.suffix}")
        return path

    @staticmethod
    def _open_doc(filepath: str) -> Document:
        """打开 Word 文档，统一入口。"""
        path = WordTool._resolve_path(filepath)
        logger.info("打开文档: %s", path)
        return Document(str(path))

    @staticmethod
    def _get_paragraph_full_text(para) -> str:
        """获取一个 paragraph 中所有 run 合并后的文本。"""
        return "".join(run.text for run in para.runs)

    @staticmethod
    def _replace_paragraph_text(para, old_text: str, new_text: str) -> int:
        """
        在单个 paragraph 中替换文本，处理跨 run 的情况。
        返回该 paragraph 中被替换的次数。
        """
        if not para.runs:
            return 0

        full_text = WordTool._get_paragraph_full_text(para)
        if old_text not in full_text:
            return 0

        count = full_text.count(old_text)
        new_full = full_text.replace(old_text, new_text)

        # 构建 run 到字符区间的映射
        # run_positions[i] = (run_index, start_char, end_char)
        run_positions: list[tuple[int, int, int]] = []
        offset = 0
        for idx, run in enumerate(para.runs):
            length = len(run.text)
            if length > 0:
                run_positions.append((idx, offset, offset + length))
            offset += length

        # 将 new_full 文本按原有 run 边界重新分配
        if len(run_positions) == 1:
            # 简单情况：只有一个 run
            para.runs[0].text = new_full
        else:
            # 多个 run：将新文本写回第一个 run，清空其余
            first_run = run_positions[0][0]
            para.runs[first_run].text = new_full
            for rp in run_positions[1:]:
                para.runs[rp[0]].text = ""

        return count

    # ------------------------------------------------------------------
    # 公开接口
    # ------------------------------------------------------------------

    @staticmethod
    def extract_text(filepath: str) -> str:
        """提取 Word 文档的全部文本内容，段落间用换行分隔。"""
        doc = WordTool._open_doc(filepath)
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        text = "\n".join(paragraphs)
        logger.info("提取文本完成，共 %d 个段落", len(paragraphs))
        return text

    @staticmethod
    def extract_paragraphs(filepath: str) -> list[str]:
        """提取所有段落，每个段落一个字符串。"""
        doc = WordTool._open_doc(filepath)
        result = [para.text for para in doc.paragraphs]
        logger.info("提取段落完成，共 %d 个段落", len(result))
        return result

    @staticmethod
    def extract_tables(filepath: str) -> list[list[list[str]]]:
        """
        提取所有表格。返回三维列表：
        [表格1[[行1[单元格1, 单元格2], 行2[...]]], 表格2[...]]
        """
        doc = WordTool._open_doc(filepath)
        tables: list[list[list[str]]] = []
        for table in doc.tables:
            table_data: list[list[str]] = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        logger.info("提取表格完成，共 %d 个表格", len(tables))
        return tables

    @staticmethod
    def replace_text(
        filepath: str,
        old_text: str,
        new_text: str,
        output_path: str | None = None,
    ) -> int:
        """
        替换文档中的文本。

        - output_path=None → 覆盖原文件
        - output_path≠None → 保存到新文件
        返回替换的次数。

        注意：python-docx 的文本可能跨 run 分段，需要处理。
        """
        doc = WordTool._open_doc(filepath)
        total = 0

        # 替换段落中的文本
        for para in doc.paragraphs:
            total += WordTool._replace_paragraph_text(para, old_text, new_text)

        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        total += WordTool._replace_paragraph_text(
                            para, old_text, new_text
                        )

        # 保存
        save_to = Path(output_path).resolve() if output_path else Path(filepath).resolve()
        doc.save(str(save_to))
        logger.info(
            "文本替换完成: 共替换 %d 处，保存至 %s", total, save_to
        )
        return total

    @staticmethod
    def replace_in_table(
        filepath: str,
        table_index: int,
        row: int,
        col: int,
        new_text: str,
        output_path: str | None = None,
    ) -> None:
        """替换指定表格的指定单元格文本。"""
        doc = WordTool._open_doc(filepath)

        if table_index < 0 or table_index >= len(doc.tables):
            raise IndexError(
                f"表格索引 {table_index} 越界，文档共有 {len(doc.tables)} 个表格"
            )

        table = doc.tables[table_index]

        if row < 0 or row >= len(table.rows):
            raise IndexError(
                f"行索引 {row} 越界，表格 {table_index} 共有 {len(table.rows)} 行"
            )

        if col < 0 or col >= len(table.rows[row].cells):
            raise IndexError(
                f"列索引 {col} 越界，表格 {table_index} 第 {row} 行共有 "
                f"{len(table.rows[row].cells)} 列"
            )

        cell = table.rows[row].cells[col]
        # 清空 cell 的原有段落，设置新文本
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""
        cell.paragraphs[0].add_run(new_text)

        save_to = (
            Path(output_path).resolve() if output_path else Path(filepath).resolve()
        )
        doc.save(str(save_to))
        logger.info(
            "表格单元格替换完成: 表格 %d, 行 %d, 列 %d → %s",
            table_index, row, col, save_to,
        )

    @staticmethod
    def create_document(filepath: str, content: list[dict]) -> None:
        """
        创建新 Word 文档。
        content 格式: [
            {"type": "heading", "text": "标题", "level": 1},
            {"type": "paragraph", "text": "段落内容"},
            {"type": "table", "data": [["A", "B"], ["1", "2"]]},
        ]
        """
        doc = Document()
        logger.info("创建新文档: %s", filepath)

        for item in content:
            item_type = item.get("type", "")

            if item_type == "heading":
                level = item.get("level", 1)
                doc.add_heading(item.get("text", ""), level=level)

            elif item_type == "paragraph":
                doc.add_paragraph(item.get("text", ""))

            elif item_type == "table":
                data = item.get("data", [])
                if not data:
                    logger.warning("表格数据为空，跳过")
                    continue

                rows_count = len(data)
                cols_count = len(data[0]) if data else 0
                table = doc.add_table(rows=rows_count, cols=cols_count)
                table.style = "Table Grid"

                for r_idx, row_data in enumerate(data):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < cols_count:
                            table.rows[r_idx].cells[c_idx].text = str(cell_text)
            else:
                logger.warning("未知的内容类型: %s，跳过", item_type)

        save_to = Path(filepath).resolve()
        save_to.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(save_to))
        logger.info("文档创建完成: %s", save_to)

    @staticmethod
    def get_info(filepath: str) -> dict:
        """
        返回文档信息:
        {"paragraphs": 段落数, "tables": 表格数, "sections": 节数}
        """
        doc = WordTool._open_doc(filepath)
        info = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
        }
        logger.info("文档信息: %s", info)
        return info
